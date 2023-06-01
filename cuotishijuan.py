#使用：用ct.xlsx内错题列表，错题在img，生成wuliti.docx错题集
from docx import Document    #导入库
from docx.shared import RGBColor    #字体颜色
from docx.shared import Cm									# cm表示
from docx.shared import Pt                                  #象素
from docx.oxml.ns import qn									# 设置中文字体用
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import openpyxl,sys
from win32com.client import Dispatch
imagePath = Path(sys.argv[0]).parent.joinpath('img')
jsq = 0
stlist = []
def just_open(filename):#代码写入的公式，程序读不出来，调用win32com.client import Dispatch 模块。运行下面代码后再进行读取
    xlApp = Dispatch("Excel.Application")
    xlApp.Visible = False
    xlBook = xlApp.Workbooks.Open(filename)
    xlBook.Save()
    xlBook.Close()
def scdoc():
    global jsq
    global stlist
    shitilist=[]
    doc = Document()    # 新建文档，如打开已有，加文件名'1.docx'
    section = doc.sections[0] # 获取section对象
    section.page_width = Cm(21) # 设置A4纸的宽度
    section.page_height = Cm(29.7) # 设置A4纸的高度
    ###纸张边距
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    ### 设置正文颜色，大小，粗体
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles['Normal'].font.size = Pt(12)
    ### 设置全文字体
    doc.styles['Normal'].font.name = u'宋体'
    ### 对于中文字体必须加这一句
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    #
    ### 添加标题，0，1，2...表示标题号大到小
    titlenr=doc.add_heading('物理练习题', 1)
    titlenr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    just_open(filename="{}".format(Path(sys.argv[0]).parent.joinpath('ct.xlsx')))
    import pandas as pd
    pabc = pd.read_excel('ct.xlsx',sheet_name='ct', index_col=0)  # 新版本的xlrd只能打开xls文件，所以需要安装老版本的1.20版本即可。sheet_name='ct',不指定索引，会自加一列索引
    # print(pabc.head(3))
    pabc.sort_values(by=["zsd1", "zsd2"], inplace=True, ascending=[True, True])  # 先对zsd1升序，再对zsd2升序排列
    # print(pabc.head(3))
    pabc.to_excel('ct1.xlsx')  # 只存操作的表，一般另存
    ctxcel = openpyxl.load_workbook('ct1.xlsx')  # 打开excel
    sh5 = ctxcel['Sheet1']
    jsq=0
    stlist=[]
    for i in range(2, sh5.max_row + 1):  # 从第2行开始，到第3行结束
        zsd1=sh5.cell(row=i, column=7).value
        zsd2=sh5.cell(row=i, column=8).value
        xxdf=sh5.cell(row=i, column=5).value
        if zsd1 != None and type(zsd2)==type(zsd2):
            zsd=zsd1+str(zsd2).replace('0','')
        else:
            zsd=zsd1
        shiti=sh5.cell(row=i, column=5).value
        if xxdf != None:
            shitilaiyuan=shiti.split('.')[0]
            # print(shiti,shitilaiyuan)
            PA = imagePath.joinpath(shiti)
            PA = Path(PA)
            # print(i,shiti,PA.is_file())
            if PA.is_file():
                # print(i, shiti)
                p3 = doc.add_paragraph()  # 加一段内容
                run = p3.add_run(u'{}.'.format(i-1))
                doc.add_picture(r"img\\{}".format(shiti), width=Cm(16.5))
                jsq=jsq+1
            else:
                stlist.append(shiti[:-4])
                # print(shiti)
            shitidaan = sh5.cell(row=i, column=6).value
            PB = imagePath.joinpath(shitidaan)
            PB =Path(PB)
            if PB.is_file():
                p3 = doc.add_paragraph()  # 加一段内容
                run = p3.add_run(u'{}.来源[{}],{} 知识点：{}'.format(i-1,shitilaiyuan,xxdf,zsd))
                doc.add_picture(r"img\\{}".format(shitidaan), width=Cm(16.5))
    p4 = doc.add_paragraph()  #
    if len(stlist) !=0:
        p4.add_run(u"{}道题导入word,没导入的{}道是{}".format(jsq,len(stlist),stlist))
    else:
        p4.add_run(u"{}道错题全部导入word".format(jsq))
    ctxcel.close()
    print("{}道题导入word：".format(jsq))
    print('没有导入的{}道是{}：'.format(len(stlist),stlist))
    doc.save('{}.docx'.format('wuli'))
if __name__ == '__main__':
    scdoc()